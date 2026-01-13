"""
Document Processor Module
Handles extraction and processing of various document formats
Includes PII detection and content sanitization
"""

import os
import re
import hashlib
from typing import Dict, List, Optional, Tuple
import logging
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document
import markdown
import chardet
from bs4 import BeautifulSoup
import textract

# NLP libraries
import spacy
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine

logger = logging.getLogger(__name__)

# Initialize spaCy and Presidio for PII detection
try:
    nlp = spacy.load("en_core_web_sm")
except:
    os.system("python -m spacy download en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# Initialize Presidio
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

class DocumentProcessor:
    """Process various document formats and extract clean text"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_doc,
            '.txt': self._extract_txt,
            '.md': self._extract_markdown,
            '.html': self._extract_html
        }
        
        # PII patterns
        self.pii_patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'(\+\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            'ssn': r'\b\d{3}-\d{2}-\d{4}\b',
            'credit_card': r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b'
        }
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from various document formats
        Returns cleaned, normalized text
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext not in self.supported_formats:
            # Try generic extraction
            return self._extract_generic(str(file_path))
        
        try:
            text = self.supported_formats[ext](str(file_path))
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Check for minimum content
            if len(text.strip()) < 100:
                raise ValueError("Document contains insufficient text content")
            
            logger.info(f"Successfully extracted {len(text)} characters from {file_path.name}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path.name}: {str(e)}")
            raise
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        text.append(page_text)
        
        except Exception as e:
            # Fallback to textract for complex PDFs
            logger.warning(f"PyPDF2 failed, trying textract: {e}")
            return self._extract_generic(file_path)
        
        return '\n'.join(text)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)
    
    def _extract_doc(self, file_path: str) -> str:
        """Extract text from DOC files (legacy Word)"""
        try:
            # Use textract for .doc files
            text = textract.process(file_path).decode('utf-8')
            return text
        except:
            # Fallback: try to convert to docx first
            logger.warning("Direct .doc extraction failed, attempting conversion")
            return self._extract_generic(file_path)
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT files with encoding detection"""
        # Detect encoding
        with open(file_path, 'rb') as file:
            raw = file.read()
            result = chardet.detect(raw)
            encoding = result['encoding'] or 'utf-8'
        
        # Read with detected encoding
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except:
            # Fallback to utf-8 with errors ignored
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
    
    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_text = file.read()
        
        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove code blocks but keep their content
        for code in soup.find_all('code'):
            code.string = f" [CODE: {code.get_text()}] "
        
        return soup.get_text()
    
    def _extract_html(self, file_path: str) -> str:
        """Extract text from HTML files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            html = file.read()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and preserve some structure
        text = soup.get_text(separator='\n')
        
        return text
    
    def _extract_generic(self, file_path: str) -> str:
        """Generic text extraction using textract"""
        try:
            text = textract.process(file_path).decode('utf-8')
            return text
        except Exception as e:
            logger.error(f"Generic extraction failed: {e}")
            # Last resort: try to read as plain text
            return self._extract_txt(file_path)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special unicode characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Fix common extraction issues
        text = re.sub(r'\.{4,}', '...', text)  # Multiple dots
        text = re.sub(r'-{3,}', '---', text)   # Multiple dashes
        text = re.sub(r'\n{3,}', '\n\n', text) # Multiple newlines
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'\[?\d+\]?\s*$', '', text, flags=re.MULTILINE)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def remove_pii(self, text: str) -> str:
        """
        Remove personally identifiable information from text
        Uses Presidio for comprehensive PII detection
        """
        try:
            # Analyze text for PII
            results = analyzer.analyze(
                text=text,
                language='en',
                entities=["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", 
                         "CREDIT_CARD", "US_SSN", "LOCATION", "DATE_TIME"]
            )
            
            # Anonymize detected PII
            if results:
                anonymized_text = anonymizer.anonymize(
                    text=text,
                    analyzer_results=results
                )
                text = anonymized_text.text
            
            # Additional pattern-based removal
            for pii_type, pattern in self.pii_patterns.items():
                text = re.sub(pattern, f'[{pii_type.upper()}_REMOVED]', text)
            
            logger.info(f"Removed {len(results)} PII entities")
            
        except Exception as e:
            logger.warning(f"PII removal failed: {e}, using basic patterns")
            # Fallback to basic pattern matching
            for pii_type, pattern in self.pii_patterns.items():
                text = re.sub(pattern, f'[{pii_type.upper()}_REMOVED]', text)
        
        return text
    
    def chunk_text(self, text: str, chunk_size: int = 1000, 
                   overlap: int = 100) -> List[str]:
        """
        Split text into overlapping chunks for processing
        Useful for large documents that exceed AI context windows
        """
        chunks = []
        sentences = text.split('. ')
        
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size > chunk_size and current_chunk:
                # Save current chunk
                chunks.append('. '.join(current_chunk) + '.')
                
                # Start new chunk with overlap
                overlap_sentences = []
                overlap_size = 0
                
                # Add sentences from the end of current chunk for overlap
                for s in reversed(current_chunk):
                    if overlap_size + len(s) <= overlap:
                        overlap_sentences.insert(0, s)
                        overlap_size += len(s)
                    else:
                        break
                
                current_chunk = overlap_sentences + [sentence]
                current_size = overlap_size + sentence_size
            else:
                current_chunk.append(sentence)
                current_size += sentence_size
        
        # Add remaining chunk
        if current_chunk:
            chunks.append('. '.join(current_chunk) + '.')
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extract document metadata"""
        file_path = Path(file_path)
        
        metadata = {
            'filename': file_path.name,
            'size': file_path.stat().st_size,
            'extension': file_path.suffix,
            'modified': file_path.stat().st_mtime
        }
        
        # PDF specific metadata
        if file_path.suffix.lower() == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf = PyPDF2.PdfReader(file)
                    if pdf.metadata:
                        metadata['title'] = pdf.metadata.get('/Title', '')
                        metadata['author'] = pdf.metadata.get('/Author', '')
                        metadata['subject'] = pdf.metadata.get('/Subject', '')
                        metadata['pages'] = len(pdf.pages)
            except:
                pass
        
        # DOCX specific metadata
        elif file_path.suffix.lower() == '.docx':
            try:
                doc = Document(file_path)
                metadata['paragraphs'] = len(doc.paragraphs)
                metadata['tables'] = len(doc.tables)
                
                # Core properties
                props = doc.core_properties
                metadata['title'] = props.title or ''
                metadata['author'] = props.author or ''
                metadata['created'] = props.created
            except:
                pass
        
        return metadata
    
    def calculate_hash(self, file_path: str) -> str:
        """Calculate file hash for deduplication"""
        sha256_hash = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        
        return sha256_hash.hexdigest()
    
    def estimate_reading_time(self, text: str, wpm: int = 200) -> int:
        """Estimate reading time in minutes"""
        word_count = len(text.split())
        return max(1, word_count // wpm)
    
    def extract_key_sentences(self, text: str, num_sentences: int = 5) -> List[str]:
        """
        Extract key sentences from text using NLP
        Useful for creating summaries or previews
        """
        doc = nlp(text[:1000000])  # Limit for spaCy
        
        # Score sentences based on various factors
        sentences = list(doc.sents)
        sentence_scores = []
        
        for sent in sentences:
            score = 0
            
            # Length factor (prefer medium-length sentences)
            if 10 < len(sent.text.split()) < 30:
                score += 1
            
            # Contains named entities
            if sent.ents:
                score += len(sent.ents)
            
            # Contains important POS tags
            for token in sent:
                if token.pos_ in ['NOUN', 'VERB']:
                    score += 0.1
            
            # Position factor (prefer early and late sentences)
            position = sentences.index(sent) / len(sentences)
            if position < 0.2 or position > 0.8:
                score += 0.5
            
            sentence_scores.append((sent.text, score))
        
        # Sort by score and return top sentences
        sentence_scores.sort(key=lambda x: x[1], reverse=True)
        
        return [s[0] for s in sentence_scores[:num_sentences]]

# Singleton instance
document_processor = DocumentProcessor()
